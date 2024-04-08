import difflib
import os
import sys
from typing import Type, Coroutine, Any
import xlrd
import numpy as np
import pandas
import numpy
import geopandas
import docx
import io
import asyncio
import win32com.client
import tempfile
from geopandas import GeoDataFrame
from pathlib import Path
from parser import *
from entities import *
from utils import *


class FileSys:

    __DEFAULT_PRJ = 'resources\\prj.txt'

    def __init__(self, input_path: Path | str = None,
                 output_path: Path | str = None,
                 prj_path: Path | str = None,
                 name_changer_path: Path | str = None,
                 template_path: Path | str = None):
        self.__set_paths(input_path, output_path)
        self.__load_prj(prj_path)
        self.__load_name_changer(name_changer_path)
        self.__load_template(template_path)

    @property
    def input_files(self) -> list[Path]:
        return [file for file in self.__input_dir.iterdir() if not file.name.startswith('~$')]

    @property
    def output_files(self) -> list[Path]:
        return list(self.__output_dir.iterdir())

    @property
    def renames(self) -> dict[str, str]:
        return self.__renames

    @property
    def template(self) -> GeoDataFrame:
        return self.__template

    @classmethod
    def read_file(cls, file: Path | str):
        file = cls.__convert_to_path(file)
        extension = file.suffix
        if extension == '.txt':
            return cls.__read_txt(file)
        elif extension == '.docx':
            return cls.__read_docx(file)
        elif extension == '.doc':
            return cls.__read_doc(file)
        elif extension == '.xls':
            return cls.__read_xls(file)
        else:
            raise NotImplemented

    @classmethod
    async def read_file_async(cls, file: Path | str):
        return cls.read_file(file)

    def save_to_shape_file(self, geodataframe: GeoDataFrame, name):
        output_dir = self.__output_dir.joinpath(name)
        if not (output_dir.exists()):
            Path.mkdir(output_dir)
        file_path = output_dir.joinpath(f'{name}.shp')
        geodataframe.to_file(str(file_path), encoding='utf-8')
        # self.__add_prj(name)

    def __set_paths(self, input_path: Path | str | None, output_path: Path | str | None) -> None:
        input_path = self.__convert_to_path(input_path)
        output_path = self.__convert_to_path(output_path)
        if input_path is None:
            self.__input_dir = self.__get_default_path('Input')
        else:
            self.__input_dir = input_path
        if output_path is None:
            self.__output_dir = self.__get_default_path('Output')
        else:
            self.__output_dir = output_path
        self.__ensure_created()

    @staticmethod
    def __convert_to_path(path: Path | str | None) -> Path | None:
        if isinstance(path, str):
            path = Path(path)
        return path

    def __get_default_path(self, path) -> Path:
        root_dir = self.__get_executable_path()
        return root_dir.joinpath(path)

    @staticmethod
    def __get_executable_path() -> Path:
        executable_path = sys.argv[0]
        abspath = os.path.abspath(executable_path)
        return Path(os.path.dirname(abspath))

    def __ensure_created(self) -> None:
        if not self.__output_dir.exists():
            Path.mkdir(self.__output_dir)
        if not self.__input_dir.exists():
            raise MissingInputDirException()

    def __load_prj(self, path: Path | str | None) -> None:
        path = self.__convert_to_path(path)
        if path is None or not path.exists():
            # self.__load_default_prj()
            self.__prj = None
            return
        self.__prj = self.__read_txt(path)[0]

    def __load_default_prj(self):
        path = self.__get_executable_path().joinpath(self.__DEFAULT_PRJ)
        if not path.exists():
            raise MissingDefaultPrjException()
        self.__prj = self.__read_txt(path)[0]

    def __load_name_changer(self, path: Path | str | None) -> None:
        path = self.__convert_to_path(path)
        if path is None:
            self.__renames = None
            return
        if not path.exists():
            self.__renames = None
            raise MissingNameChangerException
        self.__renames = NameChangerParser.parse(pandas.read_csv(path, header=None, encoding='utf-8', sep=";"))

    @staticmethod
    def __read_txt(file_path: Path) -> list[str]:
        with open(file_path, 'r') as file:
            lines = file.readlines()
        return lines

    @staticmethod
    def __read_xls(file_path: Path):
        workbook = xlrd.open_workbook(file_path)
        return workbook.sheet_by_index(0)

    @staticmethod
    def __read_docx(file_path: Path):
        return docx.Document(str(file_path)).tables

    @classmethod
    def __read_doc(cls, file_path: Path):
        return cls.__convert_doc_to_docx(file_path).tables

    @staticmethod
    def __convert_doc_to_docx(file_path: Path):
        word_app = win32com.client.Dispatch('Word.Application')
        doc = word_app.Documents.Open(str(file_path))
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_file_path = temp_file.name
        temp_file.close()
        doc.SaveAs(temp_file_path, FileFormat=16)
        doc.Close()
        word_app.Quit()
        with open(temp_file_path, 'rb') as f:
            docx_content = f.read()
        os.remove(temp_file_path)
        return docx.Document(io.BytesIO(docx_content))

    def __add_prj(self, file_name) -> None:
        if self.__prj is None:
            return
        with open(f'{file_name}.prj', 'w') as file:
            file.write(self.__prj)

    def __load_template(self, path: Path | str | None) -> None:
        path = self.__convert_to_path(path)
        if path is None or not path.exists():
            self.__template = None
            return
        self.__template = geopandas.read_file(path)
        original_template_columns = ['river', 'name', 'geometry']
        self.__template = self.__template[original_template_columns]
        self.__template = self.__template.rename(
            columns={
                'river': WaterNameProperty.dataframe_name(),
                'name': NameProperty.dataframe_name()
            }
        )


@dataclass(frozen=True)
class ParseResult:
    parser: Type[UGMSParserBase]
    objects: list[ObservationPointDTOBase]


class ConverterApp:

    def __init__(self,
                 filesys: FileSys,
                 null_symbol: str = None,
                 logger: Logger = None,
                 *parsers: Type[UGMSParserBase]):
        self.__filesys = filesys
        self.__parsers = parsers
        self.__null_symbol = null_symbol
        self.__logger = logger
        self.__create_converter()

    @property
    def null_symbol(self) -> str | None:
        return self.__null_symbol

    @null_symbol.setter
    def null_symbol(self, value: str | None) -> None:
        self.__null_symbol = value
        self.__create_converter()

    def convert(self) -> None:
        for file in self.__filesys.input_files:
            self.__convert_file(file)

    async def convert_async(self) -> None:
        tasks = list()
        for file in self.__filesys.input_files:
            tasks.append(self.__convert_file_async(file))
        await asyncio.gather(*tasks)

    def __create_converter(self) -> None:
        self.__converter = Converter(self.__filesys.template,
                                     self.__filesys.renames,
                                     self.__null_symbol,
                                     self.__logger)

    def __convert_file(self, file: Path):
        content = FileSys.read_file(file)
        parsed = self.__try_parse(content)
        result = self.__converter.convert(parsed.objects)
        self.__filesys.save_to_shape_file(result, f'{file.stem}_{parsed.parser.UGMS_CODE}')

    async def __convert_file_async(self, file: Path):
        content = await FileSys.read_file_async(file)
        parsed = await self.__try_parse_async(content)
        result = await self.__converter.convert_async(parsed.objects)
        # print(result)
        self.__filesys.save_to_shape_file(result, f'{file.stem}_{parsed.parser.UGMS_CODE}')

    def __try_parse(self, content) -> ParseResult:
        for parser in self.__parsers:
            try:
                result = parser.parse(content)
                return ParseResult(parser=parser, objects=result)
            except ParseException:
                pass
        raise UnknownFormatException()

    async def __try_parse_async(self, content) -> ParseResult:
        for parser in self.__parsers:
            try:
                result = await parser.parse_async(content)
                return ParseResult(parser=parser, objects=result)
            except ParseException:
                pass
        raise UnknownFormatException()


class Converter:

    def __init__(self,
                 template: GeoDataFrame | None,
                 renames: dict[str, str] = None,
                 null_symbol: str = None,
                 logger: Logger = None):
        self.__template = template
        self.__renames = renames
        self.__null_symbol = null_symbol
        self.__logger = logger

    async def convert_async(self, observation_stations: list[ObservationPointDTOBase]) -> GeoDataFrame:
        return self.convert(observation_stations)

    def convert(self, observation_stations: list[ObservationPointDTOBase]) -> GeoDataFrame | None:
        if self.__renames is None:
            observation_stations = self.__rename_experimental(observation_stations)
        else:
            observation_stations = self.__rename(observation_stations)
        if len(observation_stations) == 0:
            return None
        return self.__make_gdf(observation_stations)

    def __make_gdf(self, observation_stations: list[ObservationPointDTOBase]) -> GeoDataFrame:
        gdf = self.__make_template(some_point=observation_stations[0])
        names_waters = list(zip(gdf[NameProperty.dataframe_name()], gdf[WaterNameProperty.dataframe_name()]))
        for point in observation_stations:
            for index, t_point in enumerate(names_waters):
                if not (point.name in t_point and point.water_name in t_point):
                    if index == len(names_waters) - 1:
                        self.__log(point)
                    continue
                point_dict = point.to_dataframe_dict()
                for field in point_dict:
                    gdf[field][index] = point_dict[field]
                break
        self.__remove_useless_rows(gdf)
        return geopandas.GeoDataFrame(gdf, geometry='geometry')

    def __make_template(self, some_point: ObservationPointDTOBase):
        gdf = self.__template.copy()
        properties = [
            prop
            for prop in some_point.to_dataframe_dict()
            if prop != WaterNameProperty.dataframe_name() and prop != NameProperty.dataframe_name()
        ]
        for prop in properties:
            gdf[prop] = self.__null_symbol
        return gdf

    def __remove_useless_rows(self, table: GeoDataFrame) -> None:
        if self.__null_symbol is None:
            return
        levels_changes = zip(table[WaterLevelProperty.dataframe_name()],
                             table[WaterLevelChangeProperty.dataframe_name()])
        for index, point in reversed(list(enumerate(levels_changes))):
            if not (point[0] == self.__null_symbol and point[1] == self.__null_symbol):
                continue
            table.drop(table.index[index], inplace=True)

    def __rename(self, table: list[ObservationPointDTOBase]) -> list[ObservationPointDTOBase]:
        for index, point in enumerate(table):
            if point.name in self.__renames.keys():
                table[index].name = self.__renames[point.name]
            if point.water_name in self.__renames.keys():
                table[index].water_name = self.__renames[point.water_name]
        return table

    def __rename_experimental(self, table: list[ObservationPointDTOBase]) -> list[ObservationPointDTOBase]:

        waters = self.__template[WaterNameProperty.dataframe_name()].tolist()
        names = self.__template[NameProperty.dataframe_name()].tolist()

        for point_i, point in enumerate(table):

            water_match = self.__find_match(point.water_name, waters)
            name_match = self.__find_match(point.name, names)

            if name_match is not None and water_match is None:
                water_match = self.__cross_column_search(name_match, names, point.water_name, waters)
            elif name_match is None and water_match is not None:
                name_match = self.__cross_column_search(water_match, waters, point.name, names)

            if water_match is not None and name_match is not None:
                if water_match not in self.__get_values_by_indices(self.__get_value_indices(name_match, names), waters):
                    name_match = self.__cross_column_search(water_match, waters, point.name, names)

            if water_match is not None and name_match is not None:
                if name_match not in self.__get_values_by_indices(self.__get_value_indices(water_match, waters), names):
                    water_match = self.__cross_column_search(name_match, names, point.water_name, waters)

            if water_match is not None and name_match is not None:
                table[point_i].name = name_match
                table[point_i].water_name = water_match

        return table

    @classmethod
    def __cross_column_search(cls, anchor: str,
                              possible_anchors: list[str],
                              value: str,
                              possible_values: list[str]) -> str | None:
        indices = cls.__get_value_indices(anchor, possible_anchors)
        return cls.__find_match(value, cls.__get_values_by_indices(indices, possible_values))

    # Uses binary search-like algorithm
    @classmethod
    def __find_match(cls, word: str, possibilities: list[str]) -> str | None:
        threshold = 0.5
        accel = 0.25
        matches = set(difflib.get_close_matches(word, possibilities, cutoff=0.1))
        if len(matches) == 0:
            return None
        while len(matches) != 1:
            prev_match = matches
            prev_threshold = threshold
            matches = set(difflib.get_close_matches(word, matches, cutoff=threshold))
            if len(matches) == 0:
                threshold -= accel
                matches = prev_match
            else:
                threshold += accel
            accel /= 2
            if prev_threshold == threshold:
                return None
        return next(iter(matches))

    @classmethod
    def __get_value_indices(cls, value: str, in_values: list[str]) -> list[int]:
        indices = list()
        for lvi, list_value in enumerate(in_values):
            if value == list_value:
                indices.append(lvi)
        return indices

    @classmethod
    def __get_values_by_indices(cls, indices: list[int], in_values: list[str]) -> list[str]:
        values = list()
        for index in indices:
            values.append(in_values[index])
        return values

    def __log(self, point: ObservationPointDTOBase) -> None:
        if self.__logger is None:
            return
        asyncio.ensure_future(self.__logger.log(f'Не найден {point.name} {point.water_name}'))
